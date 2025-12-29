import streamlit as st
import requests, io, re, os, base64, json
from docx import Document
from typing import Dict, Any, Optional

# ======================================================
# MUST BE FIRST STREAMLIT COMMAND (KEEP ONLY ONCE)
# ======================================================
st.set_page_config(page_title="Akhila — Portfolio", layout="wide")

# ---------------------------
# CONFIG - CHANGE THESE
# ---------------------------
GITHUB_OWNER = "Akhila-A2610"
GITHUB_REPO = "portfolio"
RESUME_PATH_IN_REPO = "Akhila_A_Resume.docx"
BRANCH = "main"

# Optional local assets (stored in repo)
PROFILE_IMG = "assets/profile.jpg"
COMPANY_LOGOS = {
    "Utah State University": "assets/company_logos/usu.jfif",
    "LTIMindtree": "assets/company_logos/ltimindtree.png",
    "Tata Consultancy Services": "assets/company_logos/tcs.jfif",
}

EDU_LOGOS = {
    "Utah State University": "assets/edu_logos/USU_CS.jpg",
    "Jawaharlal Nehru Technological University": "assets/edu_logos/JNTUH.jpg",
}

LINKEDIN_USER = "akhilaa2610"


# ---------------------------
# Helpers: GitHub raw download
# ---------------------------
def download_raw_file(
    owner: str, repo: str, path: str, branch: str = "main", token: Optional[str] = None
) -> Optional[bytes]:
    raw_url = f"https://raw.githubusercontent.com/{owner}/{repo}/{branch}/{path}"
    headers = {}
    if token:
        headers["Authorization"] = f"Bearer {token}"
    r = requests.get(raw_url, headers=headers, timeout=30)
    if r.status_code == 200:
        return r.content
    return None


def download_raw_text(
    owner: str, repo: str, path: str, branch: str = "main", token: Optional[str] = None
) -> Optional[str]:
    b = download_raw_file(owner, repo, path, branch, token)
    if not b:
        return None
    try:
        return b.decode("utf-8")
    except Exception:
        return b.decode("utf-8", errors="replace")


# ---------------------------
# Parse .docx (resume)
#   NOTE: Skills removed completely.
# ---------------------------
def parse_resume_docx_bytes(docx_bytes: bytes) -> Dict[str, Any]:
    doc = Document(io.BytesIO(docx_bytes))
    para_lines = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

    parsed: Dict[str, Any] = {
        "name": "",
        "role": "Senior Data Engineer | Data Scientist",
        "contact_line": "",
        "summary": "",
        "publications": [],
        "experience": {},
        "education": [],
        "certifications": [],
    }

    if not para_lines:
        return parsed

    parsed["name"] = para_lines[0]
    if len(para_lines) > 1:
        parsed["contact_line"] = para_lines[1]

    def clean_bullet(s: str) -> str:
        return s.replace("•", "").strip()

    HEADINGS = {
        "PROFESSIONAL SUMMARY": "summary",
        "PUBLICATIONS": "publications",
        "PROFESSIONAL EXPERIENCE": "experience",
        "EDUCATION": "education",
        "CERTIFICATIONS & ACHIEVEMENTS": "certifications",
        "CERTIFICATIONS": "certifications",
        "ACHIEVEMENTS": "certifications",
        # ignore these if they exist in resume
        "TECHNICAL SKILLS": "ignore",
        "SKILLS": "ignore",
    }

    section: Optional[str] = None
    current_job: Optional[str] = None

    job_header_re = re.compile(
        r""".+,\s*.+\s+
            ((Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)
            |January|February|March|April|May|June|July|August|September|October|November|December)
            \s+\d{4}\s*[–-]\s*
            (Present|
            ((Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)
            |January|February|March|April|May|June|July|August|September|October|November|December)
            \s+\d{4})
        """,
        re.IGNORECASE | re.VERBOSE,
    )

    i = 2
    while i < len(para_lines):
        s = para_lines[i]
        s_u = s.upper().strip()

        if s_u in HEADINGS:
            section = HEADINGS[s_u]
            current_job = None
            i += 1
            continue

        if section == "ignore":
            i += 1
            continue

        if section == "summary":
            parsed["summary"] += s + "\n"
            i += 1
            continue

        if section == "publications":
            if s.strip().startswith("•"):
                parsed["publications"].append(clean_bullet(s))
            else:
                if parsed["publications"]:
                    parsed["publications"][-1] += " " + s.strip()
                else:
                    parsed["publications"].append(s.strip())
            i += 1
            continue

        if section == "experience":
            if job_header_re.search(s) and "•" not in s:
                current_job = s.strip()
                parsed["experience"][current_job] = []
                i += 1
                continue

            if current_job and s.strip().startswith("•"):
                parsed["experience"][current_job].append(clean_bullet(s))
                i += 1
                continue

            if current_job and parsed["experience"][current_job]:
                parsed["experience"][current_job][-1] += " " + s.strip()
                i += 1
                continue

            i += 1
            continue

        if section == "education":
            parsed["education"].append(s.strip())
            i += 1
            continue

        if section == "certifications":
            if s.strip().startswith("•"):
                parsed["certifications"].append(clean_bullet(s))
            else:
                if parsed["certifications"]:
                    parsed["certifications"][-1] += " " + s.strip()
                else:
                    parsed["certifications"].append(s.strip())
            i += 1
            continue

        i += 1

    parsed["summary"] = parsed["summary"].strip()
    return parsed


# ---------------------------
# Projects loader (projects.docx in repo root)
# ---------------------------
def load_projects_from_github(
    owner: str, repo: str, branch: str = "main", token: Optional[str] = None
) -> str:
    projects_path = "projects.docx"
    content_bytes = download_raw_file(owner, repo, projects_path, branch, token)
    if not content_bytes:
        return "projects.docx not found in your GitHub repo root."

    doc = Document(io.BytesIO(content_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join([f"• {p}" for p in paragraphs])


# ---------------------------
# Load resume from GitHub (Streamlit cache)
# ---------------------------
@st.cache_data(show_spinner=False)
def load_resume_from_github(
    owner: str, repo: str, path: str, branch: str = "main", token: Optional[str] = None
) -> Dict[str, Any]:
    content_bytes = download_raw_file(owner, repo, path, branch, token)
    if content_bytes is None:
        raise RuntimeError("Could not download resume file from GitHub (check file name/path).")
    return parse_resume_docx_bytes(content_bytes)


# ---------------------------
# Publications loader (publications.json)
# ---------------------------
def load_publications_from_github(
    owner: str, repo: str, branch: str = "main", token: Optional[str] = None
) -> list[dict]:
    b = download_raw_file(owner, repo, "publications.json", branch, token)
    if not b:
        return []
    try:
        return json.loads(b.decode("utf-8"))
    except Exception:
        return []


# ---------------------------
# Contact hyperlinks
# ---------------------------
def make_hyperlinked_contact(contact_text: str, linkedin_url: str, github_url: str) -> str:
    if not contact_text:
        return ""

    email_pattern = re.compile(r"[\w\.-]+@[\w\.-]+\.\w+")
    contact_text = email_pattern.sub(
        lambda m: f'<a href="mailto:{m.group(0)}" style="color:#87CEFA;text-decoration:none;">{m.group(0)}</a>',
        contact_text,
    )

    contact_text = re.sub(
        r"\bLinkedIn\b",
        f'<a href="{linkedin_url}" target="_blank" style="color:#87CEFA;text-decoration:none;">LinkedIn</a>',
        contact_text,
        flags=re.IGNORECASE,
    )
    contact_text = re.sub(
        r"\bGitHub\b",
        f'<a href="{github_url}" target="_blank" style="color:#87CEFA;text-decoration:none;">GitHub</a>',
        contact_text,
        flags=re.IGNORECASE,
    )

    contact_text = contact_text.replace("|", "&nbsp;|&nbsp;")
    return contact_text


# ---------------------------
# Logo pickers
# ---------------------------
def pick_company_key(job_header: str, logo_map: Dict[str, str]) -> Optional[str]:
    j = (job_header or "").lower()
    for k in logo_map.keys():
        if k.lower() in j:
            return k
    return None


def pick_edu_logo(edu_line: str, logo_map: Dict[str, str]) -> Optional[str]:
    s = (edu_line or "").lower()
    for key, path in logo_map.items():
        if key.lower() in s:
            return path
    return None


# ---------------------------
# UI helpers (NO cards)
# ---------------------------
def css():
    st.markdown(
        """
<style>
/* Hide Streamlit top chrome */
header { visibility: hidden; height: 0px; }
footer { visibility: hidden; height: 0px; }
[data-testid="stHeader"] { display: none; }
[data-testid="stToolbar"] { display: none; }

/* Remove default top padding */
.main .block-container { padding-top: 0rem !important; }
section.main > div { padding-top: 0rem !important; }
div.block-container { padding-top: 0rem !important; }

/* Hide accidental code rendering */
div[data-testid="stMarkdownContainer"] pre { display: none !important; }

/* Theme */
.stApp { background-color: #0b0f19; color: white; }
.muted { color: #b9c0d4; }

/* Sticky header */
.sticky {
  position: fixed;
  top: 0; left: 0;
  width: 100%;
  background: rgba(11,15,25,0.96);
  backdrop-filter: blur(10px);
  border-bottom: 2px solid rgba(135,206,250,0.75);
  z-index: 9999;
  padding: 14px 18px;
}
.header-row {
  display:flex;
  align-items:flex-start;
  justify-content:space-between;
  gap: 18px;
  max-width: 1200px;
  margin: 0 auto;
}
.id-row { display:flex; align-items:center; gap:14px; min-width: 360px; }
.avatar {
  width: 74px;
  height: 74px;
  border-radius: 50%;
  object-fit: cover;
  border: 2px solid rgba(135,206,250,0.75);
}
.nav {
  display:flex;
  flex-wrap:wrap;
  gap:10px;
  justify-content:flex-end;
  padding-top:6px;
  max-width:560px;
}
.nav a {
  background:#11a9c0;
  color:white !important;
  text-decoration:none !important;
  padding:10px 14px;
  border-radius:8px;
  font-weight:800;
  font-size:14px;
  white-space:nowrap;
}
.nav a:hover { background:#02839a; }

/* Spacer pushes content below sticky header */
.spacer { height: 155px; }
a[id] { scroll-margin-top: 175px; }

/* Section titles (Saichand-style) */
.section-title {
  font-size: 34px;
  font-weight: 900;
  margin: 22px 0 6px 0;
}
.section-sub {
  color: #b9c0d4;
  margin-bottom: 10px;
}

/* Work Experience clickable tiles */
.company-card {
  display:flex;
  flex-direction:column;
  align-items:center;
  gap: 10px;
}
.company-logo {
  background: white;
  border-radius: 12px;
  padding: 10px;
}

/* Style Streamlit buttons like top nav */
div[data-testid="stButton"] > button {
  background:#11a9c0 !important;
  color:white !important;
  border: none !important;
  border-radius:10px !important;
  padding:10px 16px !important;
  font-weight:800 !important;
  font-size:14px !important;
  width: 100% !important;
}
div[data-testid="stButton"] > button:hover {
  background:#02839a !important;
}
</style>
""",
        unsafe_allow_html=True,
    )


def render_sticky_header(name, role, contact_html, profile_img_b64=None):
    avatar_html = ""
    if profile_img_b64:
        avatar_html = f"<img class='avatar' src='data:image/jpeg;base64,{profile_img_b64}' />"

    html = (
        f'<div class="sticky">'
        f'  <div class="header-row">'
        f'    <div class="id-row">'
        f'      {avatar_html}'
        f'      <div>'
        f'        <div style="font-size:34px;font-weight:900;color:gold;line-height:1;">{name}</div>'
        f'        <div style="font-size:26px;font-weight:900;color:limegreen;line-height:1.1;">{role}</div>'
        f'        <div class="muted" style="margin-top:6px;font-size:15px;">{contact_html}</div>'
        f'      </div>'
        f'    </div>'
        f'    <div class="nav">'
        f'      <a href="#summary">Summary</a>'
        f'      <a href="#experience">Work Experience</a>'
        f'      <a href="#certs">Certifications</a>'
        f'      <a href="#publications">Publications</a>'
        f'      <a href="#projects">Projects</a>'
        f'      <a href="#education">Education</a>'
        f'      <a href="#about">About</a>'
        f'    </div>'
        f'  </div>'
        f'</div>'
        f'<div class="spacer"></div>'
    )
    st.markdown(html, unsafe_allow_html=True)


def section_anchor(anchor_id: str):
    st.markdown(f'<a id="{anchor_id}"></a>', unsafe_allow_html=True)


def section_title(title: str, subtitle: str = ""):
    st.markdown(f"<div class='section-title'>{title}</div>", unsafe_allow_html=True)
    if subtitle:
        st.markdown(f"<div class='section-sub'>{subtitle}</div>", unsafe_allow_html=True)


# ---------------------------
# Experience with logos (NO card wrapper)
# ---------------------------

def img_file_to_data_uri(path: str) -> Optional[str]:
    if not path or not os.path.exists(path):
        return None

    ext = os.path.splitext(path)[1].lower()
    mime = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
        ".gif": "image/gif",
        ".jfif": "image/jpeg",
    }.get(ext, "image/png")

    with open(path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")

    return f"data:{mime};base64,{b64}"

def render_experience_with_logos(experience: Dict[str, list], logo_map: Dict[str, str]):
    # --- helper: match company name inside job header ---
    def pick_company_key(job_header: str, logo_map: Dict[str, str]) -> Optional[str]:
        j = (job_header or "").lower()
        for k in logo_map.keys():
            if k.lower() in j:
                return k
        return None

    # --- helper: convert local image to base64 data URI (so HTML <img> works) ---
    def img_file_to_data_uri(path: str) -> Optional[str]:
        if not path or not os.path.exists(path):
            return None

        ext = os.path.splitext(path)[1].lower()
        mime = {
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".webp": "image/webp",
            ".gif": "image/gif",
            ".jfif": "image/jpeg",
        }.get(ext, "image/png")

        with open(path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode("utf-8")

        return f"data:{mime};base64,{b64}"

    if "selected_job" not in st.session_state:
        st.session_state["selected_job"] = None

    if not experience:
        return

    # Title (no card / no "click a company" text)
    st.markdown("<h2 style='margin: 6px 0 14px 0;'>Work Experience</h2>", unsafe_allow_html=True)

    # Build items (keep your current label behavior)
    items = []
    for job_header in experience.keys():
        company_key = pick_company_key(job_header, logo_map)
        logo_path = logo_map.get(company_key) if company_key else None
        label = company_key if company_key else job_header
        items.append((job_header, label, logo_path))

    cols = st.columns(min(4, len(items)))

    for idx, (job_header, label, logo_path) in enumerate(items):
        with cols[idx % len(cols)]:
            st.markdown("<div class='company-card'>", unsafe_allow_html=True)

            # Logo via base64 (prevents broken icon / white bar)
            data_uri = img_file_to_data_uri(logo_path) if logo_path else None
            if data_uri:
                st.markdown(
                    f"""
                    <div class="company-logo">
                      <img src="{data_uri}"
                           style="width:90px;height:90px;object-fit:contain;display:block;border-radius:14px;" />
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

            # Button uses your CSS (blue background, white text)
            if st.button(label, key=f"job_btn_{idx}", use_container_width=True):
                st.session_state["selected_job"] = job_header

            st.markdown("</div>", unsafe_allow_html=True)

    # Details
    selected = st.session_state.get("selected_job")
    if selected:
        bullets = experience.get(selected, [])
        with st.expander(selected, expanded=True):
            if bullets:
                st.markdown("\n".join([f"- {b}" for b in bullets]))
            else:
                st.write("No bullet points found.")
            if st.button("Close", key="close_job"):
                st.session_state["selected_job"] = None
# ---------------------------
# Main app
# ---------------------------
def main():
    css()

    with st.sidebar:
        if st.button("🔄 Refresh / Clear cache"):
            st.cache_data.clear()
            st.rerun()

    token = None
    try:
        token = st.secrets.get("GITHUB_TOKEN", None)
    except Exception:
        token = None

    resume = load_resume_from_github(GITHUB_OWNER, GITHUB_REPO, RESUME_PATH_IN_REPO, BRANCH, token)

    profile_img_b64 = None
    if os.path.exists(PROFILE_IMG):
        with open(PROFILE_IMG, "rb") as img_file:
            profile_img_b64 = base64.b64encode(img_file.read()).decode()

    linkedin_url = f"https://www.linkedin.com/in/{LINKEDIN_USER}/"
    github_url = f"https://github.com/{GITHUB_OWNER}"
    contact_html = make_hyperlinked_contact(resume.get("contact_line", ""), linkedin_url, github_url)

    render_sticky_header(
        name=resume.get("name", "Akhila A"),
        role=resume.get("role", "Senior Data Engineer | Data Scientist"),
        contact_html=contact_html,
        profile_img_b64=profile_img_b64,
    )

    # SUMMARY (no card)
    section_anchor("summary")
    section_title("Summary")
    summary_text = (resume.get("summary", "") or "").strip()
    if summary_text:
        bullets = [s.strip() for s in re.split(r"(?<=[.!?])\s+", summary_text) if s.strip()]
        st.markdown("\n".join([f"- {b}" for b in bullets]))
    else:
        st.write("No summary found in the resume.")

    # EXPERIENCE
    section_anchor("experience")
    exp = resume.get("experience", {}) or {}
    render_experience_with_logos(exp, COMPANY_LOGOS)

    # PUBLICATIONS (no card)
    section_anchor("publications")
    section_title("Publications")
    pubs = load_publications_from_github(GITHUB_OWNER, GITHUB_REPO, BRANCH, token)
    if pubs:
        for p in pubs:
            title = (p.get("title") or "").strip()
            venue = (p.get("venue") or "").strip()
            url = (p.get("url") or "").strip()
            if url:
                st.markdown(
                    f"- **[{title}]({url})**  \n  <span class='muted'>{venue}</span>",
                    unsafe_allow_html=True,
                )
            else:
                st.markdown(f"- **{title}**  \n  <span class='muted'>{venue}</span>", unsafe_allow_html=True)
    else:
        st.write("No publications found (publications.json missing or empty).")

    # CERTIFICATIONS (no card)
    section_anchor("certs")
    section_title("Certifications")
    certs = resume.get("certifications", []) or []
    if certs:
        st.markdown("\n".join([f"- {c}" for c in certs]))
    else:
        st.write("No certifications found.")

    # EDUCATION (no card, with logos)
    section_anchor("education")
    section_title("Education")
    edu_list = resume.get("education", []) or []
    if edu_list:
        for edu in edu_list:
            c1, c2 = st.columns([1, 6], vertical_alignment="center")
            with c1:
                logo_path = pick_edu_logo(edu, EDU_LOGOS)
                if logo_path and os.path.exists(logo_path):
                    st.image(logo_path, width=110)
            with c2:
                st.markdown(f"- {edu}")
    else:
        st.write("No education found.")

    # PROJECTS (no card)
    section_anchor("projects")
    section_title("Projects")
    projects_text = load_projects_from_github(GITHUB_OWNER, GITHUB_REPO, BRANCH, token)
    st.markdown(projects_text.replace("\n", "  \n"))

    # ABOUT (no card)
    section_anchor("about")
    section_title("About")
    about_txt = download_raw_text(GITHUB_OWNER, GITHUB_REPO, "aboutpage.txt", BRANCH, token)
    if about_txt:
        st.markdown(about_txt.replace("\n", "  \n"))
    else:
        st.write("aboutpage.txt not found in your GitHub repo root.")


if __name__ == "__main__":
    main()
